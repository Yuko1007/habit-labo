from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_background(cell, fill):
    """セルの背景色を設定"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._element.get_or_add_tcPr().append(shading_elm)

# Document作成
doc = Document()

# スタイル設定
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)

# タイトル
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('習慣化アプリ 要件定義書')
run.font.size = Pt(24)
run.font.bold = True
run.font.color.rgb = RGBColor(46, 117, 182)

# サブタイトル
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('起業コーチ コミュニティ向け')
run.font.size = Pt(14)
doc.add_paragraph()

# 1. プロジェクト概要
heading1 = doc.add_heading('1. プロジェクト概要', level=1)
heading1.style = 'Heading 1'
for run in heading1.runs:
    run.font.color.rgb = RGBColor(46, 117, 182)

doc.add_paragraph(
    '本アプリは、起業家向けコミュニティのメンバーが毎日の習慣をやりきる力をつけるためのサポートツールです。'
    'メンバーが設定した習慣の進捗を可視化し、「できた」を体感させることで、継続的な行動習慣を形成します。'
)

# テーブル
heading2 = doc.add_heading('1.1 アプリ名・対象・目的', level=2)
table = doc.add_table(rows=5, cols=2)
table.style = 'Light Grid Accent 1'

# ヘッダー行
header_cells = table.rows[0].cells
header_cells[0].text = 'アプリ名'
header_cells[1].text = ''
set_cell_background(header_cells[0], 'D5E8F0')
set_cell_background(header_cells[1], 'D5E8F0')

# データ行
table.rows[1].cells[0].text = 'アプリ名'
table.rows[1].cells[1].text = '習慣化トレーニングアプリ（仮称）'
set_cell_background(table.rows[1].cells[0], 'D5E8F0')

table.rows[2].cells[0].text = '対象ユーザー'
table.rows[2].cells[1].text = 'コミュニティメンバー（10名）+ 受講生（最大15名）'
set_cell_background(table.rows[2].cells[0], 'D5E8F0')

table.rows[3].cells[0].text = 'プラットフォーム'
table.rows[3].cells[1].text = 'Webブラウザ（PC・スマートフォン対応）'
set_cell_background(table.rows[3].cells[0], 'D5E8F0')

table.rows[4].cells[0].text = '主な目的'
main_goal = table.rows[4].cells[1]
main_goal.text = ''
for goal in ['メンバーの習慣化をサポート', '「できた」を体感させる', '毎日やると決めたことをやりきる力をつける', 'メンバー同士が高め合える環境を提供']:
    main_goal.add_paragraph(goal, style='List Bullet')
set_cell_background(table.rows[4].cells[0], 'D5E8F0')

doc.add_paragraph()

# 2. 機能要件
doc.add_heading('2. 機能要件', level=1)

doc.add_heading('2.1 認証機能', level=2)
for item in [
    'メール＆パスワードでログイン',
    'ログイン状態を1か月間保持（ブラウザクッキー保存）',
    '複数デバイスからのアクセスに対応',
    'パスワードリセット機能（本人確認）',
    'ユーザー登録：招待リンクから本人登録',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('2.2 習慣設定機能', level=2)
for item in [
    'メンバーが自分の習慣を自由に設定（1～2個）',
    '習慣の内容例：AI勉強、発信、LP作成など',
    '毎日、習慣の内容を変更可能',
    '習慣の追加・編集・削除が可能',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('2.3 毎日のタスク入力＆チェック機能', level=2)
for item in [
    'ホーム画面（＝自分の画面）で「今日やること」を入力',
    '入力フィールドの横に「完了」ボタンを配置',
    'チェックはその日中に実行可能、翌日以降の修正も可能',
    '設定した習慣（1～2個）ごとにチェック',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('2.4 進捗表示機能', level=2)
for item in [
    '自分の進捗：カレンダー形式で1週間・1か月単位で表示',
    '全員の進捗：カレンダー形式で全メンバーの状況を一覧表示',
    '表彰マーク：1週間ごとに自動で達成者に表示',
    '達成状況の色分け：達成（緑）/ 未達（グレー）/ 部分達成（オレンジ）など',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('2.5 管理画面（ゆうこさん専用）', level=2)
for item in [
    'デフォルト：1か月表示',
    '月選択機能：5か月分を選択して表示可能（講座が5か月間）',
    'メンバー名：ニックネーム表示',
    'メンバー削除：管理者が確認して削除',
]:
    doc.add_paragraph(item, style='List Bullet')

# 3. 画面設計
doc.add_heading('3. 画面設計', level=1)
doc.add_heading('3.1 画面一覧', level=2)

screen_table = doc.add_table(rows=7, cols=2)
screen_table.style = 'Light Grid Accent 1'

headers = screen_table.rows[0].cells
headers[0].text = '画面名'
headers[1].text = '説明'
set_cell_background(headers[0], 'D5E8F0')
set_cell_background(headers[1], 'D5E8F0')

screens = [
    ('ログイン画面', 'メール・パスワード入力、パスワードリセット'),
    ('ホーム（自分の画面）', '今日のタスク入力、完了ボタン、自分の進捗カレンダー'),
    ('全員の進捗画面', '全メンバーの進捗一覧（カレンダー形式）、表彰マーク'),
    ('習慣設定画面', '習慣の追加・編集・削除'),
    ('管理画面', '月別選択、全メンバー進捗一覧、メンバー管理（削除）'),
    ('プロフィール画面', 'ニックネーム、習慣設定、ログアウト'),
]

for i, (screen_name, description) in enumerate(screens, start=1):
    screen_table.rows[i].cells[0].text = screen_name
    screen_table.rows[i].cells[1].text = description
    set_cell_background(screen_table.rows[i].cells[0], 'D5E8F0')

# 4. 非機能要件
doc.add_heading('4. 非機能要件', level=1)

doc.add_heading('4.1 セキュリティ', level=2)
for item in [
    'メール＆パスワード認証（本人確認程度）',
    'パスワードはハッシュ化して保存',
    'ログイン状態は1か月保持',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4.2 パフォーマンス', level=2)
for item in [
    'ページ読み込み時間：3秒以内（Wi-Fi環境下）',
    'レスポンスタイム：ボタンクリック後1秒以内',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4.3 対応デバイス', level=2)
for item in [
    'PC（ブラウザ）：Chrome、Firefox、Safari、Edge',
    'スマートフォン：iOS Safari、Android Chrome（レスポンシブ対応）',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4.4 データ保持期間', level=2)
for item in [
    '5か月間（講座期間）のデータを保持',
    '期間終了後のデータは要検討',
]:
    doc.add_paragraph(item, style='List Bullet')

# 5. 補足
doc.add_heading('5. 補足', level=1)

doc.add_heading('5.1 デザイン参考', level=2)
doc.add_paragraph(
    '参考：「鬼トレ」アプリのカラースキーム（緑・ティール・オレンジ）と UI 概念を参考に、'
    'メンバーが「できた」を体感できる、わかりやすくモチベーションの上がるデザインを目指します。'
)

doc.add_heading('5.2 今後の拡張候補', level=2)
for item in [
    'Slack/LINE連携（朝のリマインダー）',
    'グループチャット・コメント機能',
    'ランキング・バッジ機能',
    'CSVエクスポート機能',
]:
    doc.add_paragraph(item, style='List Bullet')

# 保存
doc.save('習慣化アプリ_要件定義書.docx')
print('✅ 要件定義書が作成されました：習慣化アプリ_要件定義書.docx')
